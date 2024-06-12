from docx import Document

def save_results_to_docx(length, width, height, num_apartments, num_floors, volume, heat_power):
    doc = Document()
    doc.add_heading('Результаты расчетов', 0)
    doc.add_paragraph(f'Длина комнаты: {length} м')
    doc.add_paragraph(f'Ширина комнаты: {width} м')
    doc.add_paragraph(f'Высота комнаты: {height} м')
    doc.add_paragraph(f'Количество квартир: {num_apartments}')
    doc.add_paragraph(f'Количество этажей: {num_floors}')
    doc.add_paragraph(f'Объем помещения: {volume:.2f} куб. м')
    doc.add_paragraph(f'Тепловая мощность для обогрева помещения: {heat_power:.2f} кВт')
    doc.save('Результаты расчетов.docx')
